"""
Report Exporter — generates PDF, DOCX, and Excel reports with matplotlib charts.

Uses Plot-Ark brand colors (amber/stone). All labels in English.
"""

import io
import json
from datetime import datetime

# ── Matplotlib setup (non-interactive backend) ────────────────────────────────
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker

# Plot-Ark brand palette
COLORS = {
    "amber_500": "#f59e0b",
    "amber_600": "#d97706",
    "amber_300": "#fcd34d",
    "amber_50":  "#fffbeb",
    "coffee":    "#7C5C44",
    "coffee_light": "#F5EDE6",
    "coffee_mid":   "#E8D5C8",
    "stone_900": "#1c1917",
    "stone_800": "#292524",
    "stone_700": "#44403c",
    "stone_500": "#78716c",
    "stone_300": "#d6d3d1",
    "stone_200": "#e7e5e4",
    "stone_100": "#f5f5f4",
    "oat_white": "#F9F8F4",
    "oat_mid":   "#F2F1ED",
    "oat_dark":  "#E8E0D0",
    "red_500": "#ef4444",
    "red_400": "#f87171",
    "green_500": "#22c55e",
    "green_400": "#4ade80",
    "blue_500": "#3b82f6",
    "yellow_400": "#facc15",
    "bg_dark": "#1c1917",
}

BAR_COLORS = [COLORS["amber_500"], COLORS["stone_500"], COLORS["red_400"],
              COLORS["green_400"], COLORS["blue_500"], COLORS["yellow_400"]]


def _setup_chart_style():
    """Apply Plot-Ark warm oat style to matplotlib."""
    plt.rcParams.update({
        "figure.facecolor": "#FFFFFF",
        "axes.facecolor": "#FFFFFF",
        "axes.edgecolor": "none",
        "axes.labelcolor": COLORS["stone_700"],
        "axes.grid": True,
        "grid.color": COLORS["stone_200"],
        "grid.linestyle": "--",
        "grid.linewidth": 0.6,
        "grid.alpha": 0.8,
        "text.color": COLORS["stone_800"],
        "xtick.color": COLORS["stone_500"],
        "ytick.color": COLORS["stone_500"],
        "xtick.bottom": False,
        "ytick.left": False,
        "font.family": "sans-serif",
        "font.size": 10,
    })


def _fig_to_bytes(fig) -> bytes:
    """Convert a matplotlib figure to PNG bytes and close it."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="#FFFFFF")
    plt.close(fig)
    buf.seek(0)
    return buf.read()

class ReportExporter:
    """Generate PDF / DOCX / Excel from A2A analysis results."""

    def generate_charts(self, report: dict) -> dict:
        """Generate all charts as PNG bytes. Returns {chart_name: bytes}."""
        _setup_chart_style()
        charts = {}

        ba = report.get("behavior_analysis", {})
        ra = report.get("risk_assessment", {})
        co = report.get("content_optimization", {})
        cc = report.get("cohort_comparison", {})

        # 1. Module Completion Bar Chart
        modules = ba.get("module_engagement", [])
        if modules:
            fig, ax = plt.subplots(figsize=(8, 3.5))
            names = [m.get("module_name", "?")[:25] for m in modules]
            rates = [m.get("completion_rate", 0) * 100 for m in modules]
            struggles = [m.get("struggles", 0) for m in modules]

            x = range(len(names))
            bars = ax.bar(x, rates, color=COLORS["coffee_light"], edgecolor=COLORS["coffee"], linewidth=1.5)
            for bar, rate in zip(bars, rates):
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.8,
                        f"{rate:.0f}%", ha="center", va="bottom", fontsize=8, color=COLORS["stone_700"])
            short_labels = [f"M{i+1}" for i in range(len(names))]
            ax.set_xticks(x)
            ax.set_xticklabels(short_labels, rotation=0, fontsize=9)
            ax.set_ylabel("Completion Rate (%)")
            ax.set_title("Module Completion Rates", fontweight="bold", color=COLORS["stone_900"])
            plt.tight_layout()
            charts["module_completion_bar"] = _fig_to_bytes(fig)

        # 2. Engagement Trend Line
        daily = ba.get("engagement_metrics", {}).get("daily_activity", [])
        if daily:
            fig, ax = plt.subplots(figsize=(8, 3.5))
            dates = [d["date"] for d in daily]
            users = [d.get("active_users", 0) for d in daily]
            actions = [d.get("actions", 0) for d in daily]

            ax.plot(dates, users, color=COLORS["coffee"], marker="o", linewidth=2.5, label="Active Students")
            ax2 = ax.twinx()
            ax2.plot(dates, actions, color=COLORS["stone_500"], marker="s", linewidth=1.5, linestyle="--", label="Total Actions")
            ax2.set_ylabel("Actions", color=COLORS["stone_500"])
            ax.set_ylabel("Active Students", color=COLORS["coffee"])
            ax.set_title("Daily Engagement Trend", fontweight="bold")
            ax.tick_params(axis="x", rotation=45)
            lines1, labels1 = ax.get_legend_handles_labels()
            lines2, labels2 = ax2.get_legend_handles_labels()
            ax.legend(lines1 + lines2, labels1 + labels2, loc="upper right")
            plt.tight_layout()
            charts["engagement_trend_line"] = _fig_to_bytes(fig)

        # 3. Risk Distribution Pie
        risk_dist = ra.get("risk_distribution", {})
        if risk_dist and sum(risk_dist.values()) > 0:
            fig, ax = plt.subplots(figsize=(6, 6))
            labels_list = []
            sizes = []
            colors = []
            color_map = {"low": COLORS["green_500"], "medium": COLORS["yellow_400"], "high": COLORS["red_500"]}
            for level in ["low", "medium", "high"]:
                val = risk_dist.get(level, 0)
                if val > 0:
                    labels_list.append(f"{level.title()} Risk ({val})")
                    sizes.append(val)
                    colors.append(color_map[level])

            if sizes:
                wedges, texts, autotexts = ax.pie(
                    sizes, colors=colors, autopct="%1.0f%%",
                    textprops={"fontsize": 12, "fontweight": "bold"},
                    startangle=90, pctdistance=0.55,
                )
                # Make percentage text white so it's visible on colored wedges
                for at in autotexts:
                    at.set_color("white")
                # Add legend indicator so readers know what each color means
                ax.legend(
                    wedges, labels_list,
                    loc="center left", bbox_to_anchor=(1.0, 0.5),
                    fontsize=10, frameon=False,
                )
                ax.set_title("Student Risk Distribution", pad=16)
                plt.tight_layout()
                charts["risk_distribution_pie"] = _fig_to_bytes(fig)

        # 4. Verb Distribution Bar
        verbs = ba.get("verb_distribution", {})
        if verbs:
            fig, ax = plt.subplots(figsize=(8, 3.5))
            v_names = list(verbs.keys())
            v_counts = list(verbs.values())
            colors_list = [COLORS["coffee_light"] if i == 0 else COLORS["oat_mid"] for i in range(len(v_names))]
            edge_list = [COLORS["coffee"] if i == 0 else COLORS["stone_500"] for i in range(len(v_names))]
            bars = ax.barh(v_names, v_counts, color=colors_list, edgecolor=edge_list, linewidth=1.5)
            for bar, count in zip(bars, v_counts):
                ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                        str(count), va="center", fontsize=8, color=COLORS["stone_700"])
            ax.set_xlabel("Count")
            ax.set_title("Learning Activity Distribution", fontweight="bold", color=COLORS["stone_900"])
            plt.tight_layout()
            charts["verb_distribution_bar"] = _fig_to_bytes(fig)

        # 5. Cohort Comparison Grouped Bar
        groups = cc.get("groups", {})
        if groups:
            fig, ax = plt.subplots(figsize=(8, 3.5))
            g_names = [n.replace("_", " ").title() for n in groups.keys()]
            comp_rates = [groups[k].get("avg_completion", 0) * 100 for k in groups]
            struggle_rates = [groups[k].get("avg_struggle", 0) * 100 for k in groups]
            counts = [groups[k].get("count", 0) for k in groups]

            x = range(len(g_names))
            width = 0.35

            comp_bars = ax.bar([i - width/2 for i in x], comp_rates, width, color=COLORS["coffee_light"], edgecolor=COLORS["coffee"], linewidth=1.5, label="Avg Completion %")
            str_bars = ax.bar([i + width/2 for i in x], struggle_rates, width, color=COLORS["oat_mid"], edgecolor=COLORS["stone_500"], linewidth=1.5, label="Avg Struggle %")
            for bar, val in zip(list(comp_bars) + list(str_bars), comp_rates + struggle_rates):
                if val > 0:
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                            f"{val:.0f}%", ha="center", va="bottom", fontsize=7, color=COLORS["stone_700"])

            ax.set_xticks(x)
            ax.set_xticklabels([f"{n}\n(n={c})" for n, c in zip(g_names, counts)])
            ax.set_ylabel("Rate (%)")
            ax.set_title("Cohort Performance Comparison", fontweight="bold")
            ax.legend()
            plt.tight_layout()
            charts["cohort_comparison_bar"] = _fig_to_bytes(fig)

        return charts

    # ── PDF Export ─────────────────────────────────────────────────────────────

    def export_pdf(self, report: dict) -> bytes:
        """Generate a PDF report with embedded charts using ReportLab."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Image,
                Table, TableStyle, PageBreak, HRFlowable, KeepTogether
            )
        except ImportError:
            return self._fallback_text_report(report, "pdf")

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter,
                                topMargin=0.5*inch, bottomMargin=0.5*inch,
                                leftMargin=0.5*inch, rightMargin=0.5*inch)
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle("PlotArkTitle", parent=styles["Title"],
                                     textColor=colors.HexColor(COLORS["stone_900"]),
                                     fontSize=24, spaceAfter=4)
        subtitle_style = ParagraphStyle("PlotArkSubtitle", parent=styles["Heading1"],
                                        textColor=colors.HexColor(COLORS["stone_700"]),
                                        fontSize=18, spaceBefore=4, spaceAfter=4)
        meta_style = ParagraphStyle("PlotArkMeta", parent=styles["Normal"],
                                    textColor=colors.HexColor(COLORS["stone_500"]),
                                    fontSize=9, spaceAfter=12)
        h2_style = ParagraphStyle("PlotArkH2", parent=styles["Heading2"],
                                  textColor=colors.HexColor(COLORS["stone_900"]),
                                  fontSize=16, spaceAfter=2, spaceBefore=12)
        body_style = ParagraphStyle("PlotArkBody", parent=styles["Normal"],
                                    fontSize=10, spaceAfter=6)
        
        stat_card_num = ParagraphStyle("StatNum", textColor=colors.HexColor(COLORS["coffee"]), fontSize=36, fontName="Helvetica-Bold", leading=40, alignment=1, spaceAfter=4)
        stat_card_lbl = ParagraphStyle("StatLbl", textColor=colors.HexColor(COLORS["stone_500"]), fontSize=9, fontName="Helvetica-Bold", alignment=1)
        insight_style = ParagraphStyle("PlotArkInsight", parent=styles["Normal"], textColor=colors.HexColor(COLORS["stone_500"]), fontSize=9, fontName="Helvetica-Oblique", spaceBefore=4, spaceAfter=8, leftIndent=10, rightIndent=10)
        callout_style = ParagraphStyle("PlotArkCallout", textColor=colors.HexColor(COLORS["coffee"]), fontSize=48, fontName="Helvetica-Bold", leading=56, spaceAfter=8, spaceBefore=4)
        callout_label = ParagraphStyle("PlotArkCalloutLbl", textColor=colors.HexColor(COLORS["stone_500"]), fontSize=12, fontName="Helvetica-Bold", spaceAfter=16)

        def draw_header_footer(canvas, doc_obj):
            canvas.saveState()
            canvas.setFont("Helvetica-Bold", 10)
            canvas.setFillColor(colors.HexColor(COLORS["coffee"]))
            canvas.drawString(0.5*inch, letter[1] - 0.4*inch, "Plot Ark Analytics")
            canvas.setFont("Helvetica", 9)
            canvas.setFillColor(colors.HexColor(COLORS["stone_500"]))
            canvas.drawRightString(letter[0] - 0.5*inch, letter[1] - 0.4*inch, report.get("course_meta", {}).get("topic", "Untitled Course"))
            canvas.drawCentredString(letter[0]/2, 0.3*inch, f"Page {doc_obj.page} | Generated: {report.get('generated_at', datetime.now().isoformat())[:10]}")
            canvas.restoreState()

        elements = []
        charts = self.generate_charts(report)

        # ── Cover with course metadata ─────────────────────────────────────
        cm = report.get("course_meta", {})
        elements.append(Spacer(1, 1.5*inch))
        elements.append(Paragraph("Plot-Ark Analytics Report", title_style))
        elements.append(Paragraph(cm.get("topic", "Untitled Course"), subtitle_style))
        meta_parts = []
        if cm.get("level"):
            meta_parts.append(cm["level"])
        if cm.get("course_type"):
            meta_parts.append(cm["course_type"])
        if cm.get("module_count"):
            meta_parts.append(f"{cm['module_count']} modules")
        if cm.get("course_code"):
            meta_parts.append(cm["course_code"])
        meta_parts.append(f"Generated: {report.get('generated_at', datetime.now().isoformat())[:10]}")
        elements.append(Paragraph(" · ".join(meta_parts), meta_style))
        elements.append(Spacer(1, 0.3*inch))

        # Executive Summary Stat Cards
        ba = report.get("behavior_analysis", {})
        ra = report.get("risk_assessment", {})
        
        # Calculate real stats
        total_students = sum(ra.get("risk_distribution", {}).values())
        if total_students == 0:
            total_students = max([m.get("unique_students", 0) for m in ba.get("module_engagement", [])] + [0])
            
        mods = ba.get("module_engagement", [])
        avg_comp = sum(m.get("completion_rate", 0) for m in mods) / len(mods) if mods else 0
        total_at_risk = len(ra.get("at_risk_students", []))
        
        stat_data = [
            [
                [Paragraph(str(total_students), stat_card_num), Paragraph("TOTAL STUDENTS", stat_card_lbl)],
                [Paragraph(f"{avg_comp * 100:.0f}%", stat_card_num), Paragraph("AVG COMPLETION", stat_card_lbl)],
                [Paragraph(str(total_at_risk), stat_card_num), Paragraph("AT-RISK STUDENTS", stat_card_lbl)],
            ]
        ]
        t = Table(stat_data, colWidths=[2.5*inch]*3)
        t.setStyle(TableStyle([
            ("ALIGN", (0,0), (-1,-1), "CENTER"),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 0.2*inch))

        # ── Table of Contents (Index) ──────────────────────────────────────
        toc_heading = ParagraphStyle("PlotArkToC", parent=styles["Heading2"],
                                     textColor=colors.HexColor(COLORS["stone_700"]),
                                     fontSize=14, spaceAfter=8, spaceBefore=8)
        link_style = ParagraphStyle("PlotArkLink", parent=styles["Normal"],
                                    textColor=colors.HexColor(COLORS["blue_500"]),
                                    fontSize=11, spaceAfter=4)
        
        elements.append(Paragraph("Table of Contents", toc_heading))
        toc_items = [
            ("sec_behavior", "1. Behavior Analysis"),
            ("sec_risk", "2. Risk Assessment"),
            ("sec_content", "3. Content Optimization"),
            ("sec_cohort", "4. Cohort Comparison"),
            ("sec_overview", "5. Overview & Recommended Actions")
        ]
        for anchor, label in toc_items:
            elements.append(Paragraph(f'<font color="{COLORS["blue_500"]}"><a href="#{anchor}">{label}</a></font>', link_style))
        elements.append(PageBreak())

        # Section 1: Behavior Analysis
        elements.append(Paragraph('<a name="sec_behavior"/>1. Behavior Analysis', h2_style))
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#D1D5DB"), spaceBefore=2, spaceAfter=12))
        
        ba = report.get("behavior_analysis", {})
        
        # Inject Numerical Summaries
        verbs = ba.get("verb_distribution", {})
        if verbs:
            # Create a 4-column summary grid
            v_items = list(verbs.items())
            table_data = []
            row = []
            for i, (v, count) in enumerate(v_items):
                row.append(f"{count}\n{v.title()}")
                if len(row) == 4 or i == len(v_items) - 1:
                    while len(row) < 4:
                        row.append("")
                    table_data.append(row)
                    row = []
            
            t = Table(table_data, colWidths=[1.5*inch]*4)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(COLORS["coffee_light"])),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor(COLORS["stone_700"])),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("LINEBELOW", (0, 0), (-1, -2), 0.3, colors.HexColor(COLORS["oat_dark"])),
                ("LINEBEFORE", (1, 0), (-1, -1), 0.3, colors.HexColor(COLORS["oat_dark"])),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 0.2*inch))

        # Peak Activity text
        peak = ba.get("engagement_metrics", {}).get("peak_activity_hours", [])
        if peak:
            peak_str = ", ".join([f"{h}:00" for h in peak])
            elements.append(Paragraph(f"<b>Peak Activity Hours:</b> {peak_str}", body_style))
            elements.append(Spacer(1, 0.2*inch))

        if "engagement_trend_line" in charts:
            elements.append(Image(io.BytesIO(charts["engagement_trend_line"]),
                                  width=6.5*inch, height=2.8*inch))
            elements.append(Paragraph('The daily engagement trend indicates overall student activity volume versus active unique students. High peaks generally align with deadlines or new module releases.', insight_style))

        if "verb_distribution_bar" in charts:
            elements.append(Spacer(1, 0.2*inch))
            elements.append(Image(io.BytesIO(charts["verb_distribution_bar"]),
                                  width=6.5*inch, height=2.8*inch))
            elements.append(Paragraph('Action distribution reveals the dominant types of learning interactions. A high volume of passive text verbs (e.g., "viewed", "read") compared to active ones (e.g., "completed", "passed") may suggest a need for more interactive assessments.', insight_style))

        # Section 2: Risk Assessment
        elements.append(PageBreak())
        elements.append(Paragraph('<a name="sec_risk"/>2. Risk Assessment', h2_style))
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#D1D5DB"), spaceBefore=2, spaceAfter=8))
        
        # Big Number Callout + Risk Pie + At-Risk Table — all kept together on
        # the same page to avoid a large blank gap below the pie chart.
        ra = report.get("risk_assessment", {})
        high_risk_count = ra.get("risk_distribution", {}).get("high", 0)
        at_risk = ra.get("at_risk_students", [])

        risk_section_elements = []
        risk_section_elements.append(Paragraph(str(high_risk_count), callout_style))
        risk_section_elements.append(Paragraph("HIGH RISK STUDENTS NEEDING ATTENTION", callout_label))

        if "risk_distribution_pie" in charts:
            risk_section_elements.append(Image(io.BytesIO(charts["risk_distribution_pie"]),
                                               width=3.5*inch, height=3.0*inch))
            risk_section_elements.append(Paragraph(
                'The risk pie chart groups students into Low, Medium, and High risk categories '
                'based on learning progress and struggle rates.', insight_style))

        if at_risk:
            risk_section_elements.append(Paragraph(f"At-Risk Students ({len(at_risk)})", h2_style))
            table_data = [["Name", "Risk", "Score", "Key Signal"]]
            for s in at_risk[:15]:
                sig = s.get("signals", [""])[0] if s.get("signals") else ""
                table_data.append([
                    Paragraph(s["name"], body_style),
                    s["risk_level"].upper(),
                    str(s["risk_score"]),
                    Paragraph(sig, body_style)
                ])

            t = Table(table_data, colWidths=[1.5*inch, 0.8*inch, 0.6*inch, 3.5*inch], repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(COLORS["oat_dark"])),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(COLORS["stone_800"])),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("TEXTCOLOR", (0, 1), (-1, -1), colors.HexColor(COLORS["stone_700"])),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("LINEBELOW", (0, 0), (-1, 0), 0.4, colors.HexColor(COLORS["stone_300"])),
                ("LINEBELOW", (0, 1), (-1, -2), 0.3, colors.HexColor(COLORS["stone_200"])),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor(COLORS["oat_white"]), colors.HexColor(COLORS["oat_mid"])]),
            ]))
            risk_section_elements.append(t)

        elements.append(KeepTogether(risk_section_elements))

        # Section 3: Content Optimization
        elements.append(PageBreak())
        elements.append(Paragraph('<a name="sec_content"/>3. Content Optimization', h2_style))
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#D1D5DB"), spaceBefore=2, spaceAfter=8))
        
        # Inject Module Engagement Data Table
        modules = report.get("behavior_analysis", {}).get("module_engagement", [])
        if modules:
            elements.append(Paragraph("Module Engagement Summary", ParagraphStyle("PlotArkH3", parent=h2_style, fontSize=12, spaceBefore=0)))
            table_data = [["#", "Module", "Students", "Completion", "Struggle"]]
            for idx, m in enumerate(modules, 1):
                name_str = m.get("module_name", "")
                table_data.append([
                    f"M{idx}",
                    Paragraph(name_str, body_style),
                    str(m.get("unique_students", 0)),
                    f"{m.get('completion_rate', 0)*100:.0f}%",
                    f"{m.get('struggle_rate', 0)*100:.0f}%"
                ])
            t = Table(table_data, colWidths=[0.4*inch, 2.8*inch, 0.8*inch, 1.0*inch, 1.0*inch])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(COLORS["oat_dark"])),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(COLORS["stone_800"])),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("TEXTCOLOR", (0, 1), (-1, -1), colors.HexColor(COLORS["stone_700"])),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LINEBELOW", (0, 0), (-1, 0), 0.4, colors.HexColor(COLORS["stone_300"])),
                ("LINEBELOW", (0, 1), (-1, -2), 0.3, colors.HexColor(COLORS["stone_200"])),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor(COLORS["oat_white"]), colors.HexColor(COLORS["oat_mid"])]),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 0.2*inch))

        if "module_completion_bar" in charts:
            elements.append(Image(io.BytesIO(charts["module_completion_bar"]),
                                  width=6.5*inch, height=2.8*inch))
            elements.append(Paragraph('Completion rates across all modules. Significant dips commonly correlate with increased difficulty or unoptimized content delivery requiring attention.', insight_style))

        co = report.get("content_optimization", {})
        under = co.get("underperforming_content", [])
        if under:
            elements.append(Paragraph("Modules Needing Attention", h2_style))
            for m in under[:5]:
                elements.append(Paragraph(
                    f"<b>{m['module_name']}</b> — Struggle: {m['struggle_rate']:.0%}, "
                    f"Completion: {m['completion_rate']:.0%}", body_style))
                for s in m.get("suggestions", []):
                    elements.append(Paragraph(f"  → {s}", body_style))

        # Section 4: Cohort Comparison
        elements.append(PageBreak())
        elements.append(Paragraph('<a name="sec_cohort"/>4. Cohort Comparison', h2_style))
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#D1D5DB"), spaceBefore=2, spaceAfter=8))
        
        if "cohort_comparison_bar" in charts:
            elements.append(Image(io.BytesIO(charts["cohort_comparison_bar"]),
                                  width=6.5*inch, height=2.8*inch))
            elements.append(Paragraph('This bar chart segments learners by cohort risk level (High vs Avg vs Low), highlighting how initial disengagement compounds into widening struggle percentage gaps.', insight_style))

        cc = report.get("cohort_comparison", {})
        for insight in cc.get("insights", []):
            elements.append(Paragraph(f"• {insight}", body_style))

        # Section 5: Overview & Recommended Actions
        elements.append(PageBreak())
        elements.append(Paragraph('<a name="sec_overview"/>5. Overview & Recommended Actions', h2_style))
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#D1D5DB"), spaceBefore=2, spaceAfter=8))
        elements.append(Spacer(1, 0.1*inch))

        overview_lines = self._generate_overview(report)
        for line in overview_lines:
            elements.append(Paragraph(line, body_style))

        doc.build(elements, onFirstPage=draw_header_footer, onLaterPages=draw_header_footer)
        return buf.getvalue()

    # ── DOCX Export ────────────────────────────────────────────────────────────

    def export_docx(self, report: dict) -> bytes:
        """Generate DOCX report with embedded charts."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return self._fallback_text_report(report, "docx")

        COFFEE = RGBColor(0x7C, 0x5C, 0x44)
        STONE = RGBColor(0x44, 0x40, 0x3C)

        def _color_heading(heading, color=COFFEE):
            for run in heading.runs:
                run.font.color.rgb = color

        doc = Document()
        charts = self.generate_charts(report)

        # Title
        title = doc.add_heading("Plot-Ark Analytics Report", level=0)
        _color_heading(title)

        # Course subtitle
        cm = report.get("course_meta", {})
        if cm.get("topic"):
            sub = doc.add_heading(cm["topic"], level=1)
            _color_heading(sub, STONE)

        # Course metadata
        p_meta = doc.add_paragraph()
        meta_parts = []
        if cm.get("level"):
            meta_parts.append(cm["level"])
        if cm.get("course_type"):
            meta_parts.append(cm["course_type"])
        if cm.get("module_count"):
            meta_parts.append(f"{cm['module_count']} modules")
        if cm.get("course_code"):
            meta_parts.append(cm["course_code"])
        meta_parts.append(f"Generated: {report.get('generated_at', '')[:10]}")
        meta_run = p_meta.add_run(" · ".join(meta_parts))
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = STONE

        # Table of Contents — manual list
        toc_heading = doc.add_heading("Table of Contents", level=1)
        _color_heading(toc_heading)
        for item in ["1. Behavior Analysis", "2. Risk Assessment",
                     "3. Content Optimization", "4. Cohort Comparison",
                     "5. Overview & Recommended Actions"]:
            p_toc = doc.add_paragraph(item)
            p_toc.runs[0].font.color.rgb = COFFEE
            p_toc.runs[0].font.size = Pt(11)

        # Executive Summary
        ex_heading = doc.add_heading("Executive Summary", level=1)
        _color_heading(ex_heading)
        for point in report.get("executive_summary", []):
            doc.add_paragraph(f"• {point}")

        doc.add_page_break()

        # Section 1: Behavior Analysis
        _color_heading(doc.add_heading("1. Behavior Analysis", level=1))
        
        ba = report.get("behavior_analysis", {})
        verbs = ba.get("verb_distribution", {})
        if verbs:
            v_items = list(verbs.items())
            num_rows = (len(v_items) + 3) // 4
            table = doc.add_table(rows=num_rows, cols=4)
            table.style = 'Table Grid'
            for idx, (v, count) in enumerate(v_items):
                r = idx // 4
                c = idx % 4
                cell = table.cell(r, c)
                cell.text = f"{count}\n{v.title()}"
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        peak = ba.get("engagement_metrics", {}).get("peak_activity_hours", [])
        if peak:
            peak_str = ", ".join([f"{h}:00" for h in peak])
            p = doc.add_paragraph("Peak Activity Hours: ")
            p.add_run(peak_str).bold = True

        if "engagement_trend_line" in charts:
            doc.add_picture(io.BytesIO(charts["engagement_trend_line"]), width=Inches(5.5))
        if "verb_distribution_bar" in charts:
            doc.add_picture(io.BytesIO(charts["verb_distribution_bar"]), width=Inches(5))

        # Section 2: Risk Assessment
        _color_heading(doc.add_heading("2. Risk Assessment", level=1))
        if "risk_distribution_pie" in charts:
            doc.add_picture(io.BytesIO(charts["risk_distribution_pie"]), width=Inches(3.5))

        ra = report.get("risk_assessment", {})
        at_risk = ra.get("at_risk_students", [])
        if at_risk:
            _color_heading(doc.add_heading("At-Risk Students", level=2))
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Name", "Risk", "Score", "Key Signal"
            for s in at_risk[:15]:
                row = table.add_row().cells
                row[0].text = s["name"]
                row[1].text = s["risk_level"].upper()
                row[2].text = str(s["risk_score"])
                row[3].text = s.get("signals", [""])[0] if s.get("signals") else ""

        # Section 3: Content
        _color_heading(doc.add_heading("3. Content Optimization", level=1))

        modules = report.get("behavior_analysis", {}).get("module_engagement", [])
        if modules:
            _color_heading(doc.add_heading("Module Engagement Summary", level=2))
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Module", "Students", "Completion", "Struggle"
            for m in modules:
                row = table.add_row().cells
                row[0].text = m.get("module_name", "")
                row[1].text = str(m.get("unique_students", 0))
                row[2].text = f"{m.get('completion_rate', 0)*100:.0f}%"
                row[3].text = f"{m.get('struggle_rate', 0)*100:.0f}%"

        if "module_completion_bar" in charts:
            doc.add_picture(io.BytesIO(charts["module_completion_bar"]), width=Inches(5.5))

        co = report.get("content_optimization", {})
        under = co.get("underperforming_content", [])
        if under:
            _color_heading(doc.add_heading("Modules Needing Attention", level=2))
            for m in under[:5]:
                p = doc.add_paragraph()
                run = p.add_run(m["module_name"])
                run.bold = True
                p.add_run(f" — Struggle: {m['struggle_rate']:.0%}, Completion: {m['completion_rate']:.0%}")
                for s in m.get("suggestions", []):
                    doc.add_paragraph(f"  → {s}")

        # Section 4: Cohort
        _color_heading(doc.add_heading("4. Cohort Comparison", level=1))
        if "cohort_comparison_bar" in charts:
            doc.add_picture(io.BytesIO(charts["cohort_comparison_bar"]), width=Inches(5.5))
        for insight in report.get("cohort_comparison", {}).get("insights", []):
            doc.add_paragraph(f"• {insight}")

        # Section 5: Overview
        _color_heading(doc.add_heading("5. Overview & Recommended Actions", level=1))
        for line in self._generate_overview(report):
            doc.add_paragraph(line)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── Excel Export ───────────────────────────────────────────────────────────

    def export_excel(self, report: dict, course_id: int = 0) -> bytes:
        """Generate multi-sheet Excel with raw data + summaries."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
        except ImportError:
            return b"openpyxl not installed"

        from db import get_db

        wb = Workbook()
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill("solid", fgColor="44403C")  # stone_700
        red_fill = PatternFill("solid", fgColor="FCA5A5")
        yellow_fill = PatternFill("solid", fgColor="FDE68A")

        # ── Sheet 1: Raw xAPI Data ────────────────────────────────────────────
        ws1 = wb.active
        ws1.title = "Raw xAPI Data"
        raw_headers = ["Student Name", "Email", "Verb", "Object ID", "Object Name", "Timestamp", "Course"]
        for c, h in enumerate(raw_headers, 1):
            cell = ws1.cell(row=1, column=c, value=h)
            cell.font = header_font
            cell.fill = header_fill

        conn = get_db()
        if conn:
            try:
                cur = conn.cursor()
                cur.execute("""
                    SELECT actor_name, actor_email, verb, object_id, object_name, timestamp, curriculum_topic
                    FROM xapi_statements
                    WHERE object_id LIKE %s
                    ORDER BY timestamp
                """, (f"course/{course_id}/%",))
                for r_idx, row in enumerate(cur.fetchall(), 2):
                    for c_idx, val in enumerate(row, 1):
                        ws1.cell(row=r_idx, column=c_idx,
                                 value=val.isoformat() if hasattr(val, "isoformat") else val)
                cur.close()
                conn.close()
            except Exception as e:
                print(f"Excel raw data error: {e}")
                if conn:
                    conn.close()

        # ── Sheet 2: Module Summary ───────────────────────────────────────────
        ws2 = wb.create_sheet("Module Summary")
        ba = report.get("behavior_analysis", {})
        mod_headers = ["Module", "Students", "Completions", "Struggles",
                       "Completion Rate", "Struggle Rate", "Drop-off Rate"]
        for c, h in enumerate(mod_headers, 1):
            cell = ws2.cell(row=1, column=c, value=h)
            cell.font = header_font
            cell.fill = header_fill

        for r_idx, m in enumerate(ba.get("module_engagement", []), 2):
            ws2.cell(row=r_idx, column=1, value=m.get("module_name", ""))
            ws2.cell(row=r_idx, column=2, value=m.get("unique_students", 0))
            ws2.cell(row=r_idx, column=3, value=m.get("completions", 0))
            ws2.cell(row=r_idx, column=4, value=m.get("struggles", 0))
            ws2.cell(row=r_idx, column=5, value=min(m.get("completion_rate", 0), 1.0))
            ws2.cell(row=r_idx, column=6, value=m.get("struggle_rate", 0))
            ws2.cell(row=r_idx, column=7, value=m.get("drop_off_rate", 0))

        # ── Sheet 3: Student Roster ───────────────────────────────────────────
        ws3 = wb.create_sheet("Student Roster")
        stu_headers = ["Name", "Email", "Risk Level", "Risk Score",
                       "Total Actions", "Mastered", "Struggled", "Failed", "Completion Rate"]
        for c, h in enumerate(stu_headers, 1):
            cell = ws3.cell(row=1, column=c, value=h)
            cell.font = header_font
            cell.fill = header_fill

        ra = report.get("risk_assessment", {})
        for r_idx, s in enumerate(ra.get("at_risk_students", []), 2):
            ws3.cell(row=r_idx, column=1, value=s.get("name", ""))
            ws3.cell(row=r_idx, column=2, value=s.get("email", ""))
            risk_cell = ws3.cell(row=r_idx, column=3, value=s.get("risk_level", ""))
            ws3.cell(row=r_idx, column=4, value=s.get("risk_score", 0))
            stats = s.get("stats", {})
            ws3.cell(row=r_idx, column=5, value=stats.get("total_actions", 0))
            ws3.cell(row=r_idx, column=6, value=stats.get("mastered", 0))
            ws3.cell(row=r_idx, column=7, value=stats.get("struggled", 0))
            ws3.cell(row=r_idx, column=8, value=stats.get("failed", 0))
            ws3.cell(row=r_idx, column=9, value=stats.get("completion_rate", 0))

            # Conditional formatting
            if s.get("risk_level") == "high":
                risk_cell.fill = red_fill
            elif s.get("risk_level") == "medium":
                risk_cell.fill = yellow_fill

        # ── Sheet 4: Feedback Summary ─────────────────────────────────────────
        ws4 = wb.create_sheet("Feedback Summary")
        fb_headers = ["Module", "Got It", "Mostly", "Something Off", "Didn't Read"]
        for c, h in enumerate(fb_headers, 1):
            cell = ws4.cell(row=1, column=c, value=h)
            cell.font = header_font
            cell.fill = header_fill

        # Get feedback from DB
        conn = get_db()
        if conn:
            try:
                cur = conn.cursor()
                cur.execute("""
                    SELECT module_index, module_title,
                        COUNT(*) FILTER (WHERE sentiment = 'got-it') as got_it,
                        COUNT(*) FILTER (WHERE sentiment = 'mostly') as mostly,
                        COUNT(*) FILTER (WHERE sentiment = 'off') as off,
                        COUNT(*) FILTER (WHERE sentiment = 'not-read') as not_read
                    FROM student_feedback
                    WHERE course_id = %s
                    GROUP BY module_index, module_title
                    ORDER BY module_index
                """, (course_id,))
                for r_idx, row in enumerate(cur.fetchall(), 2):
                    ws4.cell(row=r_idx, column=1, value=row[1] or f"Module {row[0]}")
                    ws4.cell(row=r_idx, column=2, value=row[2])
                    ws4.cell(row=r_idx, column=3, value=row[3])
                    ws4.cell(row=r_idx, column=4, value=row[4])
                    ws4.cell(row=r_idx, column=5, value=row[5])
                cur.close()
                conn.close()
            except Exception as e:
                print(f"Excel feedback error: {e}")
                if conn:
                    conn.close()

        # Auto-width columns
        for ws in [ws1, ws2, ws3, ws4]:
            for col in ws.columns:
                max_length = 0
                for cell in col:
                    try:
                        if cell.value:
                            max_length = max(max_length, len(str(cell.value)))
                    except Exception:
                        pass
                ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def _generate_overview(self, report: dict) -> list:
        """Generate Section 5 overview with data synthesis and recommendations."""
        lines = []
        ra = report.get("risk_assessment", {})
        co = report.get("content_optimization", {})
        ba = report.get("behavior_analysis", {})
        cc = report.get("cohort_comparison", {})
        cm = report.get("course_meta", {})

        # Data Summary
        lines.append("<b>Data Summary</b>")
        total = ra.get("total_students_analyzed", 0)
        at_risk = ra.get("at_risk_students", [])
        high_risk = [s for s in at_risk if s.get("risk_level") == "high"]
        lines.append(f"• Total students analyzed: {total}")
        lines.append(f"• At-risk students: {len(at_risk)} ({len(high_risk)} high-risk)")
        under = co.get("underperforming_content", [])
        high_perf = co.get("high_performing_content", [])
        lines.append(f"• Underperforming modules: {len(under)}")
        lines.append(f"• High-performing modules: {len(high_perf)}")
        groups = cc.get("groups", {})
        dis = groups.get("disengaged", {})
        if dis.get("count", 0) > 0:
            pct = (dis["count"] / max(total, 1)) * 100
            lines.append(f"• Disengaged students: {dis['count']} ({pct:.0f}%)")
        lines.append("")

        # Recommended Improvements
        lines.append("<b>Recommended Improvements</b>")
        if len(at_risk) > 5:
            lines.append(
                f"HIGH — {len(at_risk)} students are at risk. Schedule 1-on-1 check-ins "
                f"with high-risk students and provide supplementary materials."
            )
        for m in under:
            sug = m.get("suggestions", ["Review and simplify content."])[0]
            lines.append(
                f"MEDIUM — Module \"{m['module_name']}\" has a "
                f"{m['struggle_rate']:.0%} struggle rate. {sug}"
            )
        if dis.get("count", 0) > 3:
            lines.append(
                f"MEDIUM — {dis['count']} students show disengaged behavior. "
                f"Consider implementing engagement incentives or personal outreach."
            )
        peak = ba.get("engagement_metrics", {}).get("peak_activity_hours", [])
        if any(h >= 22 or h <= 5 for h in peak):
            lines.append(
                "LOW — Unusual late-night/early-morning activity detected. "
                "May indicate different time zones or poor time management."
            )
        if len(at_risk) == 0 and len(under) == 0:
            lines.append("No critical issues found. Continue monitoring student progress.")

        return lines

    def _fallback_text_report(self, report: dict, fmt: str) -> bytes:
        """Generate plain-text fallback if libraries not installed."""
        lines = ["Plot-Ark Analytics Report", "=" * 40, ""]
        for point in report.get("executive_summary", []):
            lines.append(f"• {point}")
        lines.append("")
        lines.append("(Install reportlab/python-docx for full report)")
        return "\n".join(lines).encode("utf-8")
