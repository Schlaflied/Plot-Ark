"""
DOCX Report Exporter — generates Word document with embedded charts.

Uses Plot-Ark brand palette from chart_generator.
"""

import io

from services.chart_generator import COLORS, generate_charts


def export_docx(report: dict) -> bytes:
    """Generate DOCX report with embedded charts."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        lines = ["Plot-Ark Analytics Report", "=" * 40, ""]
        for point in report.get("executive_summary", []):
            lines.append(f"• {point}")
        lines.append("")
        lines.append("(Install python-docx for full report)")
        return "\n".join(lines).encode("utf-8")

    COFFEE = RGBColor(0x7C, 0x5C, 0x44)
    STONE = RGBColor(0x44, 0x40, 0x3C)

    def _color_heading(heading, color=COFFEE):
        for run in heading.runs:
            run.font.color.rgb = color

    doc = Document()
    charts = generate_charts(report)

    # Brand line
    p_brand = doc.add_paragraph()
    p_brand.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_brand = p_brand.add_run("PLOT ARK ANALYTICS")
    r_brand.font.size = Pt(10)
    r_brand.font.bold = True
    r_brand.font.color.rgb = COFFEE

    # Spacer paragraph
    doc.add_paragraph()

    # Report type line
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_type = p_type.add_run("Plot-Ark Analytics Report")
    r_type.font.size = Pt(12)
    r_type.font.color.rgb = STONE

    # Course name (large title)
    cm = report.get("course_meta", {})
    title = doc.add_heading(cm.get("topic", "Untitled Course"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _color_heading(title)

    # Course metadata line
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_parts = []
    if cm.get("level"):
        meta_parts.append(cm["level"])
    if cm.get("course_type"):
        meta_parts.append(cm["course_type"])
    if cm.get("module_count"):
        meta_parts.append(f"{cm['module_count']} modules")
    if cm.get("course_code"):
        meta_parts.append(cm["course_code"])
    meta_run = p_meta.add_run(" · ".join(meta_parts))
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = STONE

    # Blank line before rule
    doc.add_paragraph()

    # Horizontal rule (thin paragraph border at bottom)
    p_rule = doc.add_paragraph()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p_rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "6B7280")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Compute stats
    ra = report.get("risk_assessment", {})
    ba_cover = report.get("behavior_analysis", {})
    mods_cover = ba_cover.get("module_engagement", [])
    avg_comp_cover = sum(m.get("completion_rate", 0) for m in mods_cover) / len(mods_cover) if mods_cover else 0
    total_students_cover = sum(ra.get("risk_distribution", {}).values()) or max(
        [m.get("unique_students", 0) for m in mods_cover] + [0]
    )
    total_at_risk = len(ra.get("at_risk_students", []))
    high = ra.get("risk_distribution", {}).get("high", 0)
    medium = ra.get("risk_distribution", {}).get("medium", 0)
    generated_date = report.get("generated_at", "")[:10]

    # 4-column metadata table (no borders)
    from docx.oxml import OxmlElement as _OXE
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    # Remove all borders
    tbl_pr = tbl._tbl.get_or_add_tblPr()
    tbl_bdr = _OXE("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = _OXE(f"w:{side}")
        el.set(qn("w:val"), "none")
        tbl_bdr.append(el)
    tbl_pr.append(tbl_bdr)

    # Column widths: roughly equal
    col_widths = [1.5, 1.5, 1.5, 2.0]
    for i, w in enumerate(col_widths):
        for row in tbl.rows:
            row.cells[i].width = int(w * 914400)  # EMU: 1 inch = 914400

    labels = ["GENERATED", "TOTAL STUDENTS", "AVG COMPLETION", "AT-RISK"]
    values = [
        generated_date,
        str(total_students_cover),
        f"{avg_comp_cover * 100:.0f}%",
        f"{total_at_risk} total  ({high} high / {medium} med)",
    ]

    STONE_500 = RGBColor(0x6B, 0x72, 0x80)
    STONE_700 = RGBColor(0x37, 0x41, 0x51)

    for col_idx, (lbl, val) in enumerate(zip(labels, values)):
        # Label row
        lbl_cell = tbl.rows[0].cells[col_idx]
        lbl_p = lbl_cell.paragraphs[0]
        lbl_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lbl_r = lbl_p.add_run(lbl)
        lbl_r.font.size = Pt(8)
        lbl_r.font.bold = True
        lbl_r.font.color.rgb = STONE_500
        # Value row
        val_cell = tbl.rows[1].cells[col_idx]
        val_p = val_cell.paragraphs[0]
        val_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        val_r = val_p.add_run(val)
        val_r.font.size = Pt(11)
        val_r.font.color.rgb = STONE_700

    # Blank line after table
    doc.add_paragraph()

    # Table of Contents — manual list
    toc_heading = doc.add_heading("Table of Contents", level=1)
    _color_heading(toc_heading)
    for item in ["1. Behavior Analysis", "2. Risk Assessment",
                 "3. Content Optimization",
                 "4. Feedback Signals & Cross-Validation",
                 "5. Student Comments",
                 "6. Cohort Comparison", "7. Analysis History",
                 "8. Overview & Recommended Actions"]:
        p_toc = doc.add_paragraph(item)
        p_toc.runs[0].font.color.rgb = COFFEE
        p_toc.runs[0].font.size = Pt(11)

    # Executive Summary
    ex_heading = doc.add_heading("Executive Summary", level=1)
    _color_heading(ex_heading)
    for point in report.get("executive_summary", []):
        doc.add_paragraph(f"• {point}")

    doc.add_page_break()

    # Section 1: Behavior Analysis
    _color_heading(doc.add_heading("1. Behavior Analysis", level=1))

    ba = report.get("behavior_analysis", {})
    verbs = ba.get("verb_distribution", {})
    if verbs:
        MASTERY_VERBS = ["completed", "passed"]
        STRUGGLE_VERBS = ["struggled", "failed"]
        engagement_verbs = [v for v in verbs if v not in MASTERY_VERBS and v not in STRUGGLE_VERBS]

        total = sum(verbs.values())

        def _group_count(verb_list):
            return sum(verbs.get(v, 0) for v in verb_list)

        def _pct(count):
            return f"{(count / total * 100):.1f}%" if total > 0 else "0.0%"

        def _breakdown(verb_list):
            return ", ".join(f"{v}: {verbs[v]}" for v in verb_list if v in verbs) or "—"

        mastery_count = _group_count(MASTERY_VERBS)
        struggle_count = _group_count(STRUGGLE_VERBS)
        engagement_count = _group_count(engagement_verbs)

        p_total = doc.add_paragraph()
        total_students = report.get("risk_assessment", {}).get("total_students_analyzed", 0)
        p_total.add_run(f"Total: {total} learning events across {total_students} students").font.size = Pt(9)

        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'

        headers = ["Mastery", "Struggle", "Engagement"]
        counts = [mastery_count, struggle_count, engagement_count]
        pcts = [_pct(mastery_count), _pct(struggle_count), _pct(engagement_count)]
        breakdowns = [_breakdown(MASTERY_VERBS), _breakdown(STRUGGLE_VERBS), _breakdown(engagement_verbs)]

        # Row 0: group labels (bold)
        for col, label in enumerate(headers):
            cell = table.cell(0, col)
            cell.text = ""
            run = cell.paragraphs[0].add_run(label)
            run.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Row 1: counts (bold, large)
        for col, count in enumerate(counts):
            cell = table.cell(1, col)
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(count))
            run.bold = True
            run.font.size = Pt(18)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Row 2: percentages
        for col, pct in enumerate(pcts):
            cell = table.cell(2, col)
            cell.text = pct + " of total"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Row 3: verb breakdown
        for col, bd in enumerate(breakdowns):
            cell = table.cell(3, col)
            cell.text = bd
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    peak = ba.get("engagement_metrics", {}).get("peak_activity_hours", [])
    if peak:
        peak_str = ", ".join([f"{h}:00" for h in peak])
        p = doc.add_paragraph("Peak Activity Hours: ")
        p.add_run(peak_str).bold = True

    if "engagement_trend_line" in charts:
        doc.add_picture(io.BytesIO(charts["engagement_trend_line"]), width=Inches(5.5))
    if "verb_distribution_bar" in charts:
        doc.add_picture(io.BytesIO(charts["verb_distribution_bar"]), width=Inches(5))

    # Section 2: Risk Assessment
    _color_heading(doc.add_heading("2. Risk Assessment", level=1))
    if "risk_distribution_pie" in charts:
        doc.add_picture(io.BytesIO(charts["risk_distribution_pie"]), width=Inches(3.5))

    ra = report.get("risk_assessment", {})
    at_risk = ra.get("at_risk_students", [])
    if at_risk:
        _color_heading(doc.add_heading("At-Risk Students", level=2))
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Name", "Risk", "Score", "Key Signal"
        for s in at_risk[:15]:
            row = table.add_row().cells
            row[0].text = s["name"]
            row[1].text = s["risk_level"].upper()
            row[2].text = str(s["risk_score"])
            row[3].text = s.get("signals", [""])[0] if s.get("signals") else ""

    # Section 3: Content
    _color_heading(doc.add_heading("3. Content Optimization", level=1))

    modules = report.get("behavior_analysis", {}).get("module_engagement", [])
    if modules:
        _color_heading(doc.add_heading("Module Engagement Summary", level=2))
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Module", "Students", "Completion", "Struggle"
        for m in modules:
            row = table.add_row().cells
            row[0].text = m.get("module_name", "")
            row[1].text = str(m.get("unique_students", 0))
            row[2].text = f"{m.get('completion_rate', 0)*100:.0f}%"
            row[3].text = f"{m.get('struggle_rate', 0)*100:.0f}%"

    if "module_completion_bar" in charts:
        doc.add_picture(io.BytesIO(charts["module_completion_bar"]), width=Inches(5.5))

    co = report.get("content_optimization", {})
    under = co.get("underperforming_content", [])
    if under:
        _color_heading(doc.add_heading("Modules Needing Attention", level=2))
        for m in under[:5]:
            p = doc.add_paragraph()
            run = p.add_run(m["module_name"])
            run.bold = True
            p.add_run(f" — Struggle: {m['struggle_rate']:.0%}, Completion: {m['completion_rate']:.0%}")
            for s in m.get("suggestions", []):
                doc.add_paragraph(f"  → {s}")

    # ── Time-on-Task ──────────────────────────────────────────────────────────
    time_on_task = report.get("behavior_analysis", {}).get("time_on_task", [])
    if time_on_task:
        p = doc.add_paragraph()
        run = p.add_run("Time-on-Task Analysis")
        run.bold = True
        tt_table = doc.add_table(rows=1, cols=5)
        tt_table.style = "Table Grid"
        for i, h in enumerate(["Module", "Mean", "Median", "P90", "Outliers"]):
            tt_table.rows[0].cells[i].text = h
        for t in time_on_task:
            labels = ", ".join(f"{k}: {v}" for k, v in (t.get("outlier_labels") or {}).items())
            row = tt_table.add_row().cells
            row[0].text = f"M{t.get('module_index', 0) + 1}"
            row[1].text = f"{t.get('mean_minutes', 0)} min"
            row[2].text = f"{t.get('median_minutes', 0)} min"
            row[3].text = f"{t.get('p90_minutes', 0)} min"
            row[4].text = labels or "—"

    # Section 4: Feedback Signals & Cross-Validation
    doc.add_page_break()
    _color_heading(doc.add_heading("4. Feedback Signals & Cross-Validation", level=1))

    co_fb = report.get("content_optimization", {})
    fb_signals = co_fb.get("feedback_signals", [])

    if fb_signals:
        _color_heading(doc.add_heading("Module Feedback Distribution", level=2))
        fb_table = doc.add_table(rows=1, cols=7)
        fb_table.style = "Table Grid"
        for i, h in enumerate(["#", "Module", "Got it", "Mostly", "Confused", "Didn't read", "Skip"]):
            cell = fb_table.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for idx, fb in enumerate(fb_signals, 1):
            row = fb_table.add_row().cells
            row[0].text = f"M{idx}"
            row[1].text = fb.get("module_name", "")[:42]
            row[2].text = str(fb.get("got_it", 0))
            row[3].text = str(fb.get("mostly", 0))
            row[4].text = str(fb.get("confused", 0))
            row[5].text = str(fb.get("unread", 0))
            row[6].text = str(fb.get("skip_count", 0))

        p_legend = doc.add_paragraph()
        r_legend = p_legend.add_run("Got it = green  |  Mostly = amber  |  Confused = red  |  Didn't read = dark  |  Skip = no response")
        r_legend.font.size = Pt(8)
        r_legend.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # Cross-validation flags
        flagged = [fb for fb in fb_signals if fb.get("cross_flags")]
        if flagged:
            _color_heading(doc.add_heading("Cross-Validation Flags", level=2))
            for fb in flagged:
                for flag in fb.get("cross_flags", []):
                    p = doc.add_paragraph()
                    p.add_run(fb.get("module_name", "")[:35]).bold = True
                    p.add_run(f" — {flag}")
    else:
        doc.add_paragraph("No feedback data collected yet.")

    # Section 5: Student Comments
    doc.add_page_break()
    _color_heading(doc.add_heading("5. Student Comments", level=1))
    text_comments = co_fb.get("text_comments", [])
    if text_comments:
        p_cnt = doc.add_paragraph()
        r_cnt = p_cnt.add_run(f"{len(text_comments)} open-text responses collected across modules.")
        r_cnt.font.size = Pt(9)
        r_cnt.font.color.rgb = STONE_500
        for c in text_comments:
            mod_label = f"Module {c.get('module_index', 0) + 1}"
            p = doc.add_paragraph()
            p.add_run(f"[{mod_label}] ").bold = True
            p.add_run(f'"{c.get("comment", "")}"').italic = True
    else:
        doc.add_paragraph("No student comments collected yet.")

    # Section 6: Cohort
    doc.add_page_break()
    _color_heading(doc.add_heading("6. Cohort Comparison", level=1))
    if "cohort_comparison_bar" in charts:
        doc.add_picture(io.BytesIO(charts["cohort_comparison_bar"]), width=Inches(5.5))
    for insight in report.get("cohort_comparison", {}).get("insights", []):
        doc.add_paragraph(f"• {insight}")

    # Section 7: Analysis History
    _color_heading(doc.add_heading("7. Analysis History", level=1))
    p_hist_desc = doc.add_paragraph()
    r_hist = p_hist_desc.add_run(
        "Historical analysis snapshots showing how key metrics evolve over time. "
        "Improvements after curriculum changes indicate design effectiveness."
    )
    r_hist.font.size = Pt(9)
    r_hist.font.color.rgb = STONE_500

    try:
        from db import get_db as _get_db_docx
        _conn = _get_db_docx()
        if _conn:
            _cur = _conn.cursor()
            _cur.execute("""
                SELECT run_at, total_students, at_risk_count,
                       module_engagement_summary, noise_label
                FROM course_analysis_snapshots
                WHERE course_id = %s
                ORDER BY run_at ASC
                LIMIT 10
            """, (report.get("course_id"),))
            _rows = _cur.fetchall()
            _cur.close()
            _conn.close()

            if _rows:
                table = doc.add_table(rows=1, cols=6)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text, hdr[1].text, hdr[2].text = "Run Date", "Students", "At-Risk %"
                hdr[3].text, hdr[4].text, hdr[5].text = "Completion", "Noise", "Trend"
                prev_risk = None
                for row_data in _rows:
                    run_at, total, at_risk, mod_json, noise = row_data
                    ar_pct = round(at_risk / max(total, 1) * 100, 1)
                    avg_comp = 0
                    if mod_json:
                        import json as _hjson
                        mods_data = mod_json if isinstance(mod_json, list) else _hjson.loads(mod_json)
                        rates = [m.get("completion_rate", 0) for m in mods_data if isinstance(m, dict)]
                        avg_comp = round(sum(rates) / max(len(rates), 1) * 100, 1)
                    trend = ""
                    if prev_risk is not None:
                        delta = ar_pct - prev_risk
                        if delta < -2:
                            trend = "Improving"
                        elif delta > 2:
                            trend = "Worsening"
                        else:
                            trend = "Stable"
                    prev_risk = ar_pct
                    date_str = run_at.strftime("%Y-%m-%d %H:%M") if run_at else ""
                    row = table.add_row().cells
                    row[0].text = date_str
                    row[1].text = str(total)
                    row[2].text = f"{ar_pct}%"
                    row[3].text = f"{avg_comp}%"
                    row[4].text = noise or ""
                    row[5].text = trend
            else:
                doc.add_paragraph("No prior analysis runs found for this course.")
    except Exception as _hist_err:
        doc.add_paragraph(f"Could not load history: {_hist_err}")

    # Trend chart image (matplotlib)
    try:
        from services.chart_generator import generate_history_chart
        hist_chart_bytes = generate_history_chart(report.get("course_id"))
        if hist_chart_bytes:
            doc.add_picture(io.BytesIO(hist_chart_bytes), width=Inches(5.5))
            p_chart_note = doc.add_paragraph()
            r_note = p_chart_note.add_run(
                "Trend: at-risk % (red) vs completion rate (blue) across analysis runs. "
                "Declining at-risk after curriculum changes indicates effective optimization."
            )
            r_note.font.size = Pt(8)
            r_note.font.color.rgb = STONE_500
    except Exception:
        pass

    # Section 8: Overview
    _color_heading(doc.add_heading("8. Overview & Recommended Actions", level=1))
    from services.export_pdf import _generate_overview
    for line in _generate_overview(report):
        # Strip HTML tags for plain-text DOCX
        clean = line.replace("<b>", "").replace("</b>", "")
        doc.add_paragraph(clean)

    # AI disclaimer
    disclaimer_para = doc.add_paragraph()
    disclaimer_run = disclaimer_para.add_run("AI-generated insights. For reference only — use alongside your professional judgment. Data source: xAPI learner records.")
    disclaimer_run.font.size = Pt(7)
    disclaimer_run.font.color.rgb = RGBColor(0x9E, 0x8E, 0x7E)
    disclaimer_para.paragraph_format.space_before = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
