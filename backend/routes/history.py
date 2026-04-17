"""History CRUD, favorite toggle, and DOCX export routes."""

import io
import json
from flask import Blueprint, request, jsonify, send_file
from db import get_db

history_bp = Blueprint("history", __name__)


@history_bp.route("/api/history", methods=["GET"])
def get_history():
    conn = get_db()
    if not conn:
        return {"history": [], "error": "DB unavailable"}, 200
    try:
        cur = conn.cursor()
        cur.execute(
            "SELECT id, created_at, topic, level, course_code, course_type, module_count, is_favorite, semester "
            "FROM curricula ORDER BY is_favorite DESC, created_at DESC LIMIT 50"
        )
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return {"history": [
            {
                "id": r[0],
                "created_at": r[1].isoformat(),
                "topic": r[2],
                "level": r[3],
                "course_code": r[4] or "",
                "course_type": r[5] or "mixed",
                "module_count": r[6],
                "is_favorite": r[7] or False,
                "semester": r[8] or "",
            }
            for r in rows
        ]}
    except Exception as e:
        return {"history": [], "error": str(e)}, 200


@history_bp.route("/api/history/<int:curriculum_id>", methods=["GET"])
def get_curriculum_by_id(curriculum_id):
    conn = get_db()
    if not conn:
        return {"error": "DB unavailable"}, 503
    try:
        cur = conn.cursor()
        cur.execute(
            "SELECT topic, level, audience, course_code, course_type, modules, sources, semester "
            "FROM curricula WHERE id = %s",
            (curriculum_id,)
        )
        row = cur.fetchone()
        cur.close()
        conn.close()
        if not row:
            return {"error": "Not found"}, 404
        def _decode(val):
            if isinstance(val, str):
                try:
                    return json.loads(val)
                except Exception:
                    return []
            return val if val is not None else []
        modules = _decode(row[5])
        sources = _decode(row[6])
        for m in (modules if isinstance(modules, list) else []):
            if not isinstance(m.get("learning_objectives"), list):
                m["learning_objectives"] = []
            if not isinstance(m.get("recommended_readings"), list):
                m["recommended_readings"] = []
            if not isinstance(m.get("assignments"), list):
                m["assignments"] = []
        return {
            "topic": row[0],
            "level": row[1],
            "audience": row[2],
            "course_code": row[3] or "",
            "course_type": row[4] or "mixed",
            "modules": modules,
            "sources": sources,
            "semester": row[7] or "",
        }
    except Exception as e:
        return {"error": str(e)}, 500


@history_bp.route("/api/history/<int:curriculum_id>", methods=["DELETE"])
def delete_curriculum(curriculum_id):
    conn = get_db()
    if not conn:
        return {"error": "DB unavailable"}, 503
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM curricula WHERE id = %s", (curriculum_id,))
        conn.commit()
        cur.close()
        conn.close()
        return {"status": "deleted"}
    except Exception as e:
        return {"error": str(e)}, 500


@history_bp.route("/api/history/<int:curriculum_id>/favorite", methods=["POST"])
def toggle_favorite(curriculum_id):
    conn = get_db()
    if not conn:
        return {"error": "DB unavailable"}, 503
    try:
        cur = conn.cursor()
        cur.execute(
            "UPDATE curricula SET is_favorite = NOT is_favorite WHERE id = %s RETURNING is_favorite",
            (curriculum_id,)
        )
        row = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()
        return {"is_favorite": row[0] if row else False}
    except Exception as e:
        return {"error": str(e)}, 500


@history_bp.route("/api/curriculum/export/docx", methods=["POST"])
def export_docx():
    from docx import Document as DocxDocument

    data = request.get_json()
    doc = DocxDocument()

    citation_format = data.get("citation_format", "apa")

    def fmt_citation(title, url, fmt):
        from datetime import date
        today = date.today().strftime("%Y-%m-%d")
        try:
            from urllib.parse import urlparse
            domain = urlparse(url).netloc.replace("www.", "")
        except:
            domain = url
        if fmt == "apa":
            return f"{title}. Retrieved from {url}"
        elif fmt == "mla":
            return f'"{title}." {domain}, {url}'
        elif fmt == "chicago":
            return f'"{title}." Accessed {today}. {url}'
        return title

    LEVEL_LABELS = {
        'undergraduate-year-1': 'Undergraduate — Year 1',
        'undergraduate-year-2': 'Undergraduate — Year 2',
        'undergraduate-year-3': 'Undergraduate — Year 3',
        'undergraduate-year-4': 'Undergraduate — Year 4',
        'graduate': 'Graduate',
        'phd': 'PhD',
        'professional': 'Professional',
    }
    raw_level = data.get("level", "")
    level_label = LEVEL_LABELS.get(raw_level, raw_level)

    topic = data.get("topic", "Curriculum")
    title_para = doc.add_heading(topic, level=1)
    title_para.runs[0].bold = True

    meta_parts = [p for p in [data.get("course_code"), level_label, data.get("audience")] if p]
    if meta_parts:
        doc.add_paragraph(" / ".join(meta_parts))

    course_narrative = data.get("course_narrative", "")
    if course_narrative:
        doc.add_heading("Course Narrative", level=2)
        doc.add_paragraph(course_narrative)

    all_readings = []
    seen_urls = set()

    for mod in data.get("modules", []):
        mod_num = mod.get("module_number", "")
        mod_title = mod.get("title", "")
        heading_text = f"Module {mod_num}: {mod_title}" if mod_num else mod_title
        doc.add_heading(heading_text, level=2)

        objectives = mod.get("learning_objectives", [])
        if objectives:
            doc.add_paragraph("Learning Objectives", style="Normal").runs[0].bold = True if doc.paragraphs[-1].runs else None
            p = doc.paragraphs[-1]
            if p.runs:
                p.runs[0].bold = True
            for obj in objectives:
                doc.add_paragraph(obj, style="List Bullet")

        readings = mod.get("recommended_readings", [])
        if readings:
            rp = doc.add_paragraph("Readings")
            if rp.runs:
                rp.runs[0].bold = True
            for r in readings:
                doc.add_paragraph(r.get('title', ''), style="List Bullet")
                key = r.get('url') or r.get('title', '')
                if key and key not in seen_urls:
                    seen_urls.add(key)
                    all_readings.append(r)

        assignments = mod.get("assignments", [])
        if assignments:
            ap = doc.add_paragraph("Assessment")
            if ap.runs:
                ap.runs[0].bold = True
            for a in assignments:
                a_name = a.get("title") or a.get("name", "")
                a_desc = a.get("task_description", "")
                name_para = doc.add_paragraph(a_name)
                if name_para.runs:
                    name_para.runs[0].bold = True
                if a_desc:
                    doc.add_paragraph(a_desc)

    if all_readings:
        doc.add_heading("References", level=2)
        for i, r in enumerate(all_readings, start=1):
            text = fmt_citation(r.get('title', ''), r.get('url', ''), citation_format)
            doc.add_paragraph(f"[{i}] {text}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = topic.lower().replace(" ", "_") + "_curriculum.docx"
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )
